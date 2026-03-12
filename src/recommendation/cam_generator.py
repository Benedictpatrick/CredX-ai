"""
CAM Generator — Produces the final Credit Appraisal Memo.

Assembles all pipeline data into a professional, structured CAM report
covering the 26 core points required for mid-to-large corporate proposals.
Generates both structured data (CAMReport) and a formatted narrative.
"""
from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any

from loguru import logger

from src.models.schemas import (
    CAMReport,
    CreditDecision,
    DebateResult,
    FiveCsAssessment,
    FraudReport,
    LoanApplication,
    ResearchReport,
)
from src.utils.llm_client import llm_client


class CAMGenerator:
    """Generates a comprehensive Credit Appraisal Memo from all pipeline data."""

    async def generate(self, state: dict[str, Any]) -> CAMReport:
        app: LoanApplication | None = state.get("application")
        financials = state.get("financial_metrics", [])
        bank_summaries = state.get("bank_summaries", [])
        gst_summaries = state.get("gst_summaries", [])
        fraud: FraudReport | None = state.get("fraud_report")
        research: ResearchReport | None = state.get("research_report")
        site_visits = state.get("site_visits", [])
        interviews = state.get("management_interviews", [])
        five_cs: FiveCsAssessment | None = state.get("five_cs")
        debate: DebateResult | None = state.get("debate_result")
        decision: CreditDecision | None = state.get("decision")
        shap_data = state.get("shap_explanation")
        audit_trail = state.get("audit_trail", [])

        logger.info("CAM Generator: Assembling final Credit Appraisal Memo")

        if not app:
            raise ValueError("CAM generation requires an application")

        # Generate executive summary via LLM
        exec_summary = await self._generate_executive_summary(
            app, five_cs, fraud, research, decision
        )

        # Generate risk narrative via LLM
        risk_narrative = await self._generate_risk_narrative(
            fraud, research, five_cs, decision, shap_data
        )

        cam = CAMReport(
            generated_at=datetime.utcnow(),
            application=app,
            financial_metrics=financials,
            bank_summaries=bank_summaries,
            gst_summaries=gst_summaries,
            fraud_report=fraud,
            research_report=research,
            site_visits=site_visits,
            management_interviews=interviews,
            five_cs=five_cs,
            debate_result=debate,
            decision=decision,
            shap_explanation=shap_data,
            executive_summary=exec_summary,
            risk_narrative=risk_narrative,
            audit_trail=audit_trail,
        )

        logger.info(f"CAM Generated: {cam.cam_id}")
        return cam

    async def _generate_executive_summary(
        self, app, five_cs, fraud, research, decision
    ) -> str:
        """Generate the executive summary section of the CAM."""
        composite_score = f"{five_cs.composite_score:.3f}" if five_cs else "N/A"
        character_score = f"{five_cs.character_score:.2f}" if five_cs else "N/A"
        capacity_score = f"{five_cs.capacity_score:.2f}" if five_cs else "N/A"
        capital_score = f"{five_cs.capital_score:.2f}" if five_cs else "N/A"
        collateral_score = f"{five_cs.collateral_score:.2f}" if five_cs else "N/A"
        conditions_score = f"{five_cs.conditions_score:.2f}" if five_cs else "N/A"
        fraud_score = f"{fraud.overall_fraud_score:.2f}" if fraud else "N/A"
        litigation_score = f"{research.litigation_score:.2f}" if research else "N/A"
        news_sentiment = f"{research.news_sentiment:.2f}" if research else "N/A"

        prompt = f"""Write a professional executive summary for a Credit Appraisal Memo (CAM).
This summary will be presented to the bank's credit committee.

Borrower: {app.company.name if app else 'N/A'}
CIN: {app.company.cin if app else 'N/A'}
Sector: {app.company.sector if app else 'N/A'}
Requested Amount: ₹{app.requested_amount_cr if app else 'N/A'} Cr
Loan Purpose: {app.loan_purpose if app else 'N/A'}
Tenure: {app.loan_tenure_months if app else 60} months

Five Cs Composite Score: {composite_score}
- Character: {character_score}
- Capacity: {capacity_score}
- Capital: {capital_score}
- Collateral: {collateral_score}
- Conditions: {conditions_score}

Fraud Assessment: {fraud.severity.value if fraud else 'N/A'} (Score: {fraud_score})
Litigation Risk: {litigation_score}
News Sentiment: {news_sentiment}

Decision: {decision.decision.value if decision else 'N/A'}
Risk Grade: {decision.risk_grade.value if decision else 'N/A'}
Approved Amount: ₹{decision.approved_amount_cr or 'N/A'} Cr
Interest Rate: {decision.interest_rate_pct or 'N/A'}%

Write 3-4 concise paragraphs. Use professional banking language.
Include the recommendation (approve/reject/conditional) with key reasoning.
No markdown formatting. Plain text only."""

        try:
            return (await llm_client.generate(prompt)).strip()
        except Exception as e:
            logger.warning(f"Executive summary generation failed: {e}")
            return self._fallback_summary(app, decision, five_cs)

    async def _generate_risk_narrative(
        self, fraud, research, five_cs, decision, shap_data
    ) -> str:
        """Generate the risk assessment and mitigation narrative."""
        fraud_text = "No fraud analysis available."
        if fraud:
            signals_text = "; ".join(
                f"{s.signal_type}: {s.description}" for s in fraud.signals[:5]
            )
            fraud_text = (
                f"Fraud Score: {fraud.overall_fraud_score:.2f} ({fraud.severity.value}). "
                f"Circular Trading: {fraud.circular_trading_detected}. "
                f"Revenue Inflation: {fraud.revenue_inflation_detected}. "
                f"Signals: {signals_text}"
            )

        research_text = "No research available."
        if research:
            lit_cases = "; ".join(
                f"{r.case_type.value}: {r.summary}" for r in research.litigation_records[:3]
            )
            research_text = (
                f"Litigation Score: {research.litigation_score:.2f} "
                f"({len(research.litigation_records)} cases). "
                f"Cases: {lit_cases}. "
                f"Sector: {research.sector_outlook}"
            )

        shap_text = ""
        if shap_data:
            top_neg = shap_data.get("top_negative", [])
            if top_neg:
                shap_text = "Key risk drivers (SHAP): " + ", ".join(
                    f"{k} ({v:+.3f})" for k, v in top_neg
                )

        prompt = f"""Write the Risk Assessment & Mitigation section of a CAM.

{fraud_text}

{research_text}

{shap_text}

Decision: {decision.decision.value if decision else 'N/A'}
Rejection reasons: {'; '.join(decision.rejection_reasons) if decision and decision.rejection_reasons else 'None'}
Conditions: {'; '.join(decision.conditions) if decision and decision.conditions else 'None'}

Write 2-3 paragraphs covering:
1. Key risks identified by the system
2. SHAP-based reasoning for decision drivers
3. Proposed mitigants or conditions
Use professional, specific language. No markdown."""

        try:
            return (await llm_client.generate(prompt)).strip()
        except Exception as e:
            logger.warning(f"Risk narrative generation failed: {e}")
            return "Risk narrative could not be generated. See structured data."

    def _fallback_summary(self, app, decision, five_cs) -> str:
        """Deterministic fallback if LLM is unavailable."""
        parts = []
        if app:
            parts.append(
                f"Credit Appraisal Memo for {app.company.name} ({app.company.cin}). "
                f"Loan request: ₹{app.requested_amount_cr} Cr for {app.loan_purpose}."
            )
        if decision:
            parts.append(
                f"Recommendation: {decision.decision.value}. "
                f"Risk Grade: {decision.risk_grade.value}."
            )
            if decision.approved_amount_cr:
                parts.append(
                    f"Approved Amount: ₹{decision.approved_amount_cr} Cr at "
                    f"{decision.interest_rate_pct}% interest."
                )
            if decision.rejection_reasons:
                parts.append(
                    f"Rejection: {'; '.join(decision.rejection_reasons)}"
                )
        if five_cs:
            parts.append(
                f"Five Cs Composite: {five_cs.composite_score:.3f}."
            )
        return " ".join(parts) or "Executive summary unavailable."

    def export_docx(self, cam: CAMReport) -> bytes:
        """Export a CAM report to a DOCX document for judge/demo workflows."""
        from docx import Document  # type: ignore[reportMissingImports]

        doc = Document()
        doc.add_heading("Credit Appraisal Memo", level=0)

        company_name = cam.application.company.name if cam.application else "Unknown Borrower"
        doc.add_paragraph(f"Borrower: {company_name}")
        doc.add_paragraph(f"CAM ID: {cam.cam_id}")
        doc.add_paragraph(f"Generated At: {cam.generated_at.isoformat()}")

        decision = cam.decision
        if decision:
            doc.add_heading("Recommendation", level=1)
            doc.add_paragraph(f"Decision: {decision.decision.value}")
            doc.add_paragraph(f"Risk Grade: {decision.risk_grade.value}")
            doc.add_paragraph(f"Approved Amount: ₹{decision.approved_amount_cr or 'N/A'} Cr")
            doc.add_paragraph(f"Interest Rate: {decision.interest_rate_pct or 'N/A'}%")
            doc.add_paragraph(f"Risk Premium: {decision.risk_premium_pct:.2f}%")

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(cam.executive_summary or "No executive summary available.")

        doc.add_heading("Risk Narrative", level=1)
        doc.add_paragraph(cam.risk_narrative or "No risk narrative available.")

        if cam.five_cs:
            doc.add_heading("Five Cs Assessment", level=1)
            five_cs_rows = [
                ("Character", cam.five_cs.character_score, cam.five_cs.character_rationale),
                ("Capacity", cam.five_cs.capacity_score, cam.five_cs.capacity_rationale),
                ("Capital", cam.five_cs.capital_score, cam.five_cs.capital_rationale),
                ("Collateral", cam.five_cs.collateral_score, cam.five_cs.collateral_rationale),
                ("Conditions", cam.five_cs.conditions_score, cam.five_cs.conditions_rationale),
            ]
            for label, score, rationale in five_cs_rows:
                doc.add_paragraph(f"{label}: {score:.2f} — {rationale}")

        if cam.fraud_report and cam.fraud_report.signals:
            doc.add_heading("Fraud Signals", level=1)
            for signal in cam.fraud_report.signals[:8]:
                doc.add_paragraph(
                    f"{signal.signal_type}: {signal.description}",
                    style="List Bullet",
                )

        if cam.research_report:
            doc.add_heading("Secondary Research", level=1)
            doc.add_paragraph(f"Sector Outlook: {cam.research_report.sector_outlook}")
            doc.add_paragraph(f"Litigation Score: {cam.research_report.litigation_score:.2f}")
            for alert in cam.research_report.regulatory_alerts[:5]:
                doc.add_paragraph(alert, style="List Bullet")

        if decision and decision.conditions:
            doc.add_heading("Conditions Precedent", level=1)
            for condition in decision.conditions:
                doc.add_paragraph(condition, style="List Bullet")

        if decision and decision.rejection_reasons:
            doc.add_heading("Rejection Reasons", level=1)
            for reason in decision.rejection_reasons:
                doc.add_paragraph(reason, style="List Bullet")

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
